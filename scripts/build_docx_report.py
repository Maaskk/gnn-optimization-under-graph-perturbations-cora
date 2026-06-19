from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
FIGURES = ROOT / "results" / "figures"
ASSETS = REPORTS / "assets"
OUT_REPO = REPORTS / "GNN_Final_Report_2026.docx"
OUT_DOWNLOADS = Path("/Users/oussamaashad/Downloads/GNN_Final_Report_2026.docx")

RED = RGBColor(0xD7, 0x19, 0x20)
DARK_RED = RGBColor(0x9F, 0x11, 0x18)
INK = RGBColor(0x11, 0x16, 0x1F)
MUTED = RGBColor(0x5D, 0x66, 0x75)
LIGHT_RED = "FFF0F1"
LIGHT_GRAY = "F4F6F8"
LINE = "D9DEE7"
WHITE = "FFFFFF"


def read_csv_rows(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = LINE, size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_cell_margins(table, top=90, start=120, bottom=90, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run("Page ")
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def paragraph_bottom_border(paragraph, color="D71920", size="12") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)


def set_column_widths(table, widths) -> None:
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=False):
            cell.width = width


def add_para(doc, text="", size=11, color=INK, bold=False, italic=False, align=None, before=0, after=6, line=1.1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        set_run_font(
            run,
            size={1: 16, 2: 13, 3: 12}.get(level, 11),
            color=RED if level <= 2 else DARK_RED,
            bold=True,
        )
    return p


def add_callout(doc, title, body, fill=LIGHT_RED):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_cell_margins(table, top=120, bottom=120, start=160, end=160)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color="F0B8BC", size="6")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=DARK_RED, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.12
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_data_table(doc, headers, rows, widths=None, header_fill="D71920", font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_cell_margins(table)
    if widths:
        set_column_widths(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        set_cell_border(cell, color=header_fill, size="6")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(header))
        set_run_font(r, size=font_size, color=RGBColor(0xFF, 0xFF, 0xFF), bold=True)

    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for col_idx, value in enumerate(row):
            cell = cells[col_idx]
            set_cell_shading(cell, LIGHT_GRAY if row_idx % 2 else WHITE)
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=INK, bold=(col_idx == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_figure(doc, path: Path, caption: str, width=6.2):
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_run_font(r, size=9, color=MUTED, italic=True)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RED if name != "Heading 3" else DARK_RED

    header = section.header.paragraphs[0]
    header.text = "GNN Optimization under Graph Perturbations - Cora"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    for run in footer.runs:
        set_run_font(run, size=8.5, color=MUTED)


def add_cover(doc: Document) -> None:
    add_para(doc, "Master DSBD & IA - Module Algorithmes d'Optimisation", size=12, color=MUTED, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, "Projet 13 - Option 4: Robustesse des GNN face au bruit dans le graphe", size=12, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=42)
    kicker = add_para(doc, "FINAL REPORT", size=10, color=RED, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
    paragraph_bottom_border(kicker, color="D71920", size="10")
    add_para(
        doc,
        "Optimization of Graph Neural Networks under Graph Perturbations",
        size=27,
        color=INK,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=8,
        line=1.0,
    )
    add_para(
        doc,
        "Adam vs AdamW vs RMSProp vs AdaGrad vs SGD on the Cora Citation Network",
        size=14,
        color=RED,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=26,
    )
    add_callout(
        doc,
        "Main result",
        "Adam is the most reliable optimizer in the fixed protocol: 79.9% clean accuracy and 70.2% mean accuracy across all clean and perturbed experiment rows.",
    )

    metadata = [
        ("Dataset", "Cora Citation Network"),
        ("Model", "2-layer Graph Convolutional Network"),
        ("Optimizers", "Adam, AdamW, RMSProp, AdaGrad, SGD"),
        ("Perturbations", "Feature noise, edge removal, fake edge addition"),
        ("Prepared by", "Oussama Ashad and Mouhcine"),
        ("Date de remise", "20 juin 2026"),
        ("Repository", "github.com/Maaskk/gnn-optimization-under-graph-perturbations-cora"),
        ("Dashboard", "maaskk.github.io/gnn-optimization-under-graph-perturbations-cora/"),
    ]
    rows = [[label, value] for label, value in metadata]
    add_data_table(doc, ["Field", "Value"], rows, widths=[Inches(1.55), Inches(4.75)], header_fill="11161F", font_size=9.2)
    doc.add_page_break()


def add_static_toc(doc: Document) -> None:
    add_para(doc, "Table des matieres / Table of Contents", size=20, color=INK, bold=True, after=12)
    paragraph_bottom_border(doc.paragraphs[-1], color="D71920", size="10")
    entries = [
        ("1", "Abstract and key findings"),
        ("2", "Introduction"),
        ("3", "Related work"),
        ("4", "Methodology"),
        ("5", "Results"),
        ("6", "Discussion"),
        ("7", "Limitations and validity"),
        ("8", "Conclusion and recommendations"),
        ("9", "Bibliography"),
        ("10", "Annexes"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_cell_margins(table, top=100, bottom=100, start=120, end=120)
    widths = [Inches(0.55), Inches(4.9), Inches(0.95)]
    set_column_widths(table, widths)
    for cell, header in zip(table.rows[0].cells, ["#", "Section", "Focus"], strict=False):
        set_cell_shading(cell, "11161F")
        set_cell_border(cell, color="11161F")
        p = cell.paragraphs[0]
        r = p.add_run(header)
        set_run_font(r, size=9.5, color=RGBColor(255, 255, 255), bold=True)
    for num, section in entries:
        cells = table.add_row().cells
        values = [num, section, "Report"]
        for i, value in enumerate(values):
            set_cell_border(cells[i])
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            set_run_font(r, size=10.5, color=INK, bold=(i == 0))
    add_para(
        doc,
        "The contents page is static so it renders correctly in headless PDF/PNG export. The document also uses real Word heading styles for navigation.",
        size=9.5,
        color=MUTED,
        italic=True,
        after=0,
    )
    doc.add_page_break()


def fmt(value, digits=3):
    return f"{float(value):.{digits}f}"


def build_report() -> Document:
    doc = Document()
    configure_document(doc)
    doc.core_properties.author = "Oussama Ashad and Mouhcine"
    doc.core_properties.title = "Optimization of Graph Neural Networks under Graph Perturbations"
    doc.core_properties.subject = "Project 13 - Robustness of GNNs under graph noise"

    clean = read_csv_rows(ROOT / "results" / "clean_optimizer_results.csv")
    summary = read_csv_rows(ROOT / "results" / "final_optimizer_summary.csv")
    dataset = read_csv_rows(ROOT / "results" / "cora_dataset_summary.csv")[0]

    add_cover(doc)
    add_static_toc(doc)

    add_heading(doc, "1. Abstract and key findings", 1)
    add_para(
        doc,
        "This project studies how optimizer choice affects the stability and robustness of a Graph Convolutional Network trained on the Cora citation network. Five optimizers are compared under a fixed protocol: Adam, AdamW, RMSProp, AdaGrad, and SGD. The same GCN architecture, data split, seed, learning rate, weight decay, dropout, and epoch budget are used for every run.",
        line=1.18,
    )
    add_callout(
        doc,
        "Answer to the research question",
        "Adam is the best overall choice for this fixed GCN/Cora protocol. It ties AdamW on the clean graph, leads the aggregate ranking across all 65 runs, and remains strongest under edge removal and fake edge addition.",
    )
    key_rows = [
        ["Best clean accuracy", "Adam / AdamW", "0.799"],
        ["Best aggregate accuracy", "Adam", "0.702"],
        ["Best aggregate macro F1", "Adam", "0.694"],
        ["Hardest perturbation", "Feature noise", "All adaptive optimizers fall near 0.52-0.54 at 30% noise"],
        ["Weakest optimizer in fixed protocol", "SGD", "Near random-guess baseline"],
    ]
    add_data_table(doc, ["Finding", "Result", "Evidence"], key_rows, widths=[Inches(1.8), Inches(1.55), Inches(2.95)], header_fill="D71920", font_size=9)

    add_heading(doc, "2. Introduction", 1)
    add_para(
        doc,
        "Graph Neural Networks learn from both node attributes and graph topology. In citation networks such as Cora, each paper is represented as a node, citation links define graph edges, and bag-of-words features describe document content. Because GCN message passing aggregates information through neighborhoods, graph perturbations can affect both convergence and final classification quality.",
    )
    add_para(
        doc,
        "The objective is to compare optimizer behavior, not to tune each optimizer to its own best possible configuration. This is why all experiments use a common fixed protocol. The central question is: which optimizer best preserves node-classification performance when features or edges are corrupted?",
    )

    add_heading(doc, "3. Related work", 1)
    add_para(doc, "Kipf and Welling introduced the two-layer GCN model for semi-supervised classification on citation networks and reported strong Cora performance with a small labeled training set [1]. GraphSAGE and GAT later expanded the GNN family through inductive aggregation and attention mechanisms [2, 3].")
    add_para(doc, "For optimization, Adam combines momentum with adaptive per-parameter learning rates [4], AdamW decouples weight decay from Adam's adaptive update [5], RMSProp uses a moving average of squared gradients, AdaGrad accumulates squared gradients, and SGD provides the non-adaptive baseline. Prior robustness work shows that GNNs are sensitive to structure and feature perturbations [7, 8].")

    add_heading(doc, "4. Methodology", 1)
    add_heading(doc, "4.1 Dataset", 2)
    dataset_rows = [
        ["Nodes", f"{int(dataset['num_nodes']):,}", "Citation graph nodes"],
        ["Directed edges", f"{int(dataset['num_edges']):,}", "Cora edge_index entries"],
        ["Node features", f"{int(dataset['num_features']):,}", "Bag-of-words feature vector"],
        ["Classes", f"{int(dataset['num_classes'])}", "Paper topic labels"],
        ["Training nodes", "140", "Planetoid split"],
        ["Validation nodes", "500", "Planetoid split"],
        ["Test nodes", "1,000", "Planetoid split"],
    ]
    add_data_table(doc, ["Property", "Value", "Note"], dataset_rows, widths=[Inches(1.65), Inches(1.25), Inches(3.4)], header_fill="11161F", font_size=9)
    add_heading(doc, "4.2 Model and fixed protocol", 2)
    protocol_rows = [
        ["Model", "2-layer GCN", "PyTorch Geometric"],
        ["Hidden channels", "16", "GCNConv(1433 -> 16 -> 7)"],
        ["Dropout", "0.5", "Training only"],
        ["Learning rate", "0.01", "Same for all optimizers"],
        ["Weight decay", "0.0005", "Same for all optimizers"],
        ["Epochs", "200", "Fixed budget"],
        ["Seed", "42", "All experiments"],
        ["Perturbation severity", "5%, 10%, 20%, 30%", "Per perturbation type"],
    ]
    add_data_table(doc, ["Parameter", "Value", "Detail"], protocol_rows, widths=[Inches(1.65), Inches(1.55), Inches(3.1)], header_fill="11161F", font_size=9)
    add_heading(doc, "4.3 Perturbation design", 2)
    add_para(doc, "Three perturbation families are tested. Feature noise adds Gaussian noise N(0, severity) to node features. Edge removal deletes a random fraction of citation edges. Fake edge addition injects random non-existing citation links. Each optimizer is evaluated on one clean condition plus twelve perturbed conditions, for 65 total experiment rows.")

    add_heading(doc, "5. Results", 1)
    add_heading(doc, "5.1 Accuracy benchmark: is 79.9% too low?", 2)
    add_para(doc, "No. A clean accuracy of 79.9% is reasonable for a simple two-layer GCN under a single fixed optimizer-comparison protocol. The original GCN paper reports 81.5% on Cora [1], but that number is not directly identical to this experiment because it comes from its own setup, averaging, and training choices. Our goal is a fair optimizer comparison under one fixed protocol, not maximum leaderboard tuning.")
    add_callout(doc, "Interpretation", "The 70.2% aggregate accuracy is lower because it averages clean and perturbed graphs. It should not be compared directly with a clean-only Cora benchmark.")

    add_heading(doc, "5.2 Clean graph performance", 2)
    clean_rows = []
    for row in clean:
        clean_rows.append([row["optimizer"], fmt(row["test_accuracy"]), fmt(row["macro_f1"]), fmt(row["final_loss"]), f"{float(row['training_time_seconds']):.1f}s"])
    add_data_table(doc, ["Optimizer", "Accuracy", "Macro F1", "Final loss", "Time"], clean_rows, widths=[Inches(1.35), Inches(1.1), Inches(1.1), Inches(1.25), Inches(1.2)], font_size=9)
    add_figure(doc, FIGURES / "clean_accuracy_macro_f1.png", "Figure 1 - Clean Cora test accuracy and macro F1-score.", width=5.9)
    add_para(doc, "Adam and AdamW tie at 0.799 clean test accuracy. RMSProp follows very closely at 0.795. AdaGrad learns but remains lower, while SGD stays near random-guess behavior under the fixed learning rate.")

    add_heading(doc, "5.3 Aggregate performance", 2)
    add_figure(doc, ASSETS / "aggregate_accuracy_f1.png", "Figure 2 - Aggregate mean accuracy and macro F1 across all 65 runs.", width=6.2)
    pivot = {
        row["optimizer"]: {}
        for row in summary
    }
    for row in summary:
        pivot[row["optimizer"]][row["perturbation_type"]] = row["mean_test_accuracy"]
    order = ["Adam", "RMSProp", "AdamW", "AdaGrad", "SGD"]
    agg_rows = []
    for opt in order:
        agg_rows.append([
            opt,
            fmt(pivot[opt]["clean"]),
            fmt(pivot[opt]["edge_removal"]),
            fmt(pivot[opt]["fake_edge_addition"]),
            fmt(pivot[opt]["feature_noise"]),
        ])
    doc.add_page_break()
    add_data_table(doc, ["Optimizer", "Clean", "Edge removal", "Fake edges", "Feature noise"], agg_rows, widths=[Inches(1.25), Inches(1.0), Inches(1.3), Inches(1.2), Inches(1.3)], font_size=8.7)

    add_heading(doc, "5.4 Robustness by perturbation type", 2)
    add_figure(doc, FIGURES / "feature_noise_test_accuracy.png", "Figure 3 - Feature noise causes the largest accuracy collapse.", width=5.9)
    add_para(doc, "Feature noise is the hardest condition: by 30% severity, Adam, AdamW, RMSProp, and AdaGrad all converge around the 0.516-0.533 accuracy band. This suggests the disturbance challenges the feature representation and the two-layer GCN itself, not only the optimizer.")
    add_figure(doc, FIGURES / "edge_removal_accuracy.png", "Figure 4 - Edge removal is tolerated best by adaptive optimizers.", width=5.9)
    add_para(doc, "Edge removal produces the smallest drop. Adam, AdamW, and RMSProp remain close to their clean scores even at 30% edge deletion, showing that the Cora graph retains enough redundant neighborhood information within this severity range.")
    add_figure(doc, FIGURES / "fake_edge_accuracy.png", "Figure 5 - Fake edge addition causes a monotonic structural degradation.", width=5.9)
    add_para(doc, "Fake edges are more damaging than edge removal because they add misleading neighborhoods. Adam has the smallest degradation among the adaptive optimizers, while AdaGrad falls more sharply under injected links.")

    add_heading(doc, "5.5 Training dynamics", 2)
    add_figure(doc, FIGURES / "clean_loss_convergence.png", "Figure 6 - Training and validation loss over 200 epochs on the clean graph.", width=6.2)
    add_para(doc, "AdamW reaches the lowest training loss, but the wider train/validation gap indicates tighter fitting to the small 140-node training set. RMSProp reaches a strong validation plateau quickly, while SGD does not escape its near-random initial behavior.")

    add_heading(doc, "5.6 Accuracy drop matrix", 2)
    add_figure(doc, ASSETS / "accuracy_drop_heatmap.png", "Figure 7 - Accuracy drop from each optimizer's own clean baseline.", width=6.2)
    add_para(doc, "The heatmap makes the robustness pattern explicit. Feature noise generates the largest red block across adaptive optimizers. Edge removal stays pale, indicating small losses. Fake edges form an intermediate degradation pattern that grows with severity.")

    add_heading(doc, "6. Discussion", 1)
    add_para(doc, "Adam is the most balanced optimizer in this study. It does not always converge fastest, but it combines strong clean accuracy, best aggregate robustness, and reliable behavior under structural corruption. RMSProp is a strong second and is especially attractive when fast convergence to a stable plateau matters.")
    add_para(doc, "AdamW matches Adam on the clean graph, but its lower training loss does not translate into better robustness. In this fixed protocol, AdamW falls behind Adam under fake edge addition and has a wider train/validation loss gap. AdaGrad is conservative and stable under edge removal but has a lower accuracy ceiling. SGD is included as a baseline and demonstrates why adaptive optimization is important for this GCN/Cora setup.")

    add_heading(doc, "7. Limitations and validity", 1)
    limitation_rows = [
        ["Single seed", "All experiments use seed 42; multiple seeds would give confidence intervals."],
        ["Fixed learning rate", "The protocol is fair for comparison but may disadvantage SGD, which often needs tuning or momentum."],
        ["One dataset", "Cora is standard but small; PubMed and CiteSeer would test scale and transfer."],
        ["One architecture", "The conclusions apply to this 2-layer GCN, not necessarily GAT or GraphSAGE."],
        ["Random perturbations", "The perturbations are random, not adversarial attacks."],
    ]
    add_data_table(doc, ["Limitation", "Meaning"], limitation_rows, widths=[Inches(1.65), Inches(4.65)], header_fill="11161F", font_size=9)
    add_para(doc, "These limitations do not invalidate the project. They define the interpretation: this is a controlled optimizer comparison under graph noise, not a universal claim about every GNN or every citation dataset.")

    add_heading(doc, "8. Conclusion and recommendations", 1)
    add_para(doc, "The project concludes that Adam is the recommended optimizer for this GCN on Cora when robustness under graph perturbations is part of the evaluation. RMSProp is a close alternative, especially for rapid convergence. AdamW is excellent on clean Cora but less robust to fake-edge structural noise in this run. AdaGrad is cautious and stable but generally lower. SGD performs poorly under the fixed protocol.")
    add_callout(doc, "Final recommendation", "Use Adam as the default optimizer for the final project demonstration, and explain that its advantage is not only clean accuracy but also stable performance across feature and structural perturbations.")

    add_heading(doc, "9. Bibliography", 1)
    refs = [
        "[1] Kipf, T. N., & Welling, M. (2017). Semi-Supervised Classification with Graph Convolutional Networks. ICLR 2017. arxiv.org/abs/1609.02907",
        "[2] Hamilton, W., Ying, Z., & Leskovec, J. (2017). Inductive Representation Learning on Large Graphs. NeurIPS 2017. arxiv.org/abs/1706.02216",
        "[3] Velickovic, P., et al. (2018). Graph Attention Networks. ICLR 2018. arxiv.org/abs/1710.10903",
        "[4] Kingma, D. P., & Ba, J. (2015). Adam: A Method for Stochastic Optimization. ICLR 2015. arxiv.org/abs/1412.6980",
        "[5] Loshchilov, I., & Hutter, F. (2019). Decoupled Weight Decay Regularization. ICLR 2019. arxiv.org/abs/1711.05101",
        "[6] Hinton, G., Srivastava, N., & Swersky, K. (2012). RMSProp. Neural Networks for Machine Learning, Lecture 6e.",
        "[7] Zugner, D., Akbarnejad, A., & Gunnemann, S. (2018). Adversarial Attacks on Neural Networks for Graph Data. KDD 2018. arxiv.org/abs/1805.07984",
        "[8] Fey, M., & Lenssen, J. E. (2019). Fast Graph Representation Learning with PyTorch Geometric. ICLR Workshop. arxiv.org/abs/1903.02428",
    ]
    for ref in refs:
        add_para(doc, ref, size=9.5, after=4, line=1.05)

    add_heading(doc, "10. Annexes", 1)
    annex_rows = [
        ["Repository", "https://github.com/Maaskk/gnn-optimization-under-graph-perturbations-cora"],
        ["Interactive dashboard", "https://maaskk.github.io/gnn-optimization-under-graph-perturbations-cora/"],
        ["Full result table", "results/all_results.csv"],
        ["Clean results", "results/clean_optimizer_results.csv"],
        ["Feature noise results", "results/feature_noise_results.csv"],
        ["Edge removal results", "results/edge_removal_results.csv"],
        ["Fake edge results", "results/fake_edge_addition_results.csv"],
        ["Loss history", "results/clean_loss_history.csv"],
        ["Frontend source", "docs/index.html, docs/assets/app.js, docs/assets/styles.css"],
    ]
    add_data_table(doc, ["Artifact", "Location"], annex_rows, widths=[Inches(1.8), Inches(4.5)], header_fill="11161F", font_size=8.8)

    return doc


def main() -> None:
    REPORTS.mkdir(parents=True, exist_ok=True)
    doc = build_report()
    doc.save(OUT_REPO)
    doc.save(OUT_DOWNLOADS)
    print(f"Wrote {OUT_REPO}")
    print(f"Wrote {OUT_DOWNLOADS}")


if __name__ == "__main__":
    main()
